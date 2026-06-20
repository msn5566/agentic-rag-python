#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts the Technical Design Document from Markdown to Word format
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink styling
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0563C1")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)


def setup_document_styles(doc):
    """Setup custom styles for the document."""
    styles = doc.styles
    
    # Create custom heading styles if they don't exist
    try:
        # Title style
        title_style = styles['Title']
        title_style.font.size = Pt(24)
        title_style.font.bold = True
    except KeyError:
        pass
    
    try:
        # Code style
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    except ValueError:
        code_style = styles['Code']
    
    try:
        # Table style
        table_style = styles.add_style('Table Text', WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.size = Pt(10)
    except ValueError:
        table_style = styles['Table Text']
    
    return doc


def parse_markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to DOCX format."""
    
    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new document
    doc = Document()
    doc = setup_document_styles(doc)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_content = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_block_content:
                    code_text = '\n'.join(code_block_content)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Code'
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # End of table
            if table_rows:
                create_table_from_rows(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('# ').strip()
            
            if level == 1:
                p = doc.add_heading(text, level=1)
            elif level == 2:
                p = doc.add_heading(text, level=2)
            elif level == 3:
                p = doc.add_heading(text, level=3)
            else:
                p = doc.add_heading(text, level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle nested bullets
            indent_level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            # Skip table separator lines
            if re.match(r'^[\|\-\s:]+$', line):
                i += 1
                continue
                
            p = doc.add_paragraph()
            
            # Handle bold and italic formatting
            text = line.strip()
            
            # Simple bold formatting
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Handle inline code
                    code_parts = re.split(r'(`.*?`)', part)
                    for code_part in code_parts:
                        if code_part.startswith('`') and code_part.endswith('`'):
                            run = p.add_run(code_part[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                        else:
                            p.add_run(code_part)
        
        # Handle empty lines
        else:
            if not in_table:
                doc.add_paragraph()
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_rows:
        create_table_from_rows(doc, table_rows)
    
    # Save the document
    doc.save(docx_file_path)
    print(f"Successfully converted {md_file_path} to {docx_file_path}")


def create_table_from_rows(doc, table_rows):
    """Create a table from markdown table rows."""
    if len(table_rows) < 2:
        return
    
    # Parse header row
    header_row = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]
    
    # Skip separator row and get data rows
    data_rows = []
    for row in table_rows[2:]:  # Skip header and separator
        if row.strip():
            cells = [cell.strip() for cell in row.split('|')[1:-1]]
            data_rows.append(cells)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(header_row))
    table.style = 'Light Grid Accent 1'
    
    # Add header
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(header_row):
        if i < len(header_cells):
            header_cells[i].text = header_text
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Add data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                # Handle checkmarks and crosses
                cell_text = cell_text.replace('✅', '☑').replace('❌', '☒')
                row_cells[i].text = cell_text


def main():
    """Main function to convert the Technical Design Document."""
    
    # File paths
    md_file = "documents/SYSTEM_PRESENTATION.md"
    docx_file = "SYSTEM_PRESENTATION.docx"
    
    # Check if markdown file exists
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found!")
        return
    
    try:
        # Convert markdown to DOCX
        parse_markdown_to_docx(md_file, docx_file)
        print(f"\n✅ Conversion completed successfully!")
        print(f"📄 Output file: {docx_file}")
        print(f"📍 Location: {os.path.abspath(docx_file)}")
        
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
